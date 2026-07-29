from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image, KeepTogether

ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX, OUT_PDF, PUBLIC = ROOT/'output'/'docx', ROOT/'output'/'pdf', ROOT/'public'/'podcast-transcripts'
for folder in (OUT_DOCX, OUT_PDF, PUBLIC): folder.mkdir(parents=True, exist_ok=True)
LOGO = ROOT/'public'/'pcmo-logo.png'
NAVY, RED, SLATE = '0B3764', 'DC2626', '475569'

EPISODES = [
 ('episode-01-forecast-judgement','Why Reliable Forecasts Require Honest Professional Judgement','Project Controls',
  ['A forecast is a decision tool, not a promise manufactured for comfort.','Progress evidence, remaining work, assumptions, dependencies, uncertainty, and risk must be connected.','Forecast confidence improves when teams can report difficult facts early without distortion.','Leaders should respond to adverse information with disciplined action rather than pressure to preserve an unsupported date.'],
  ['Define the data date and evidence rules.','Separate actual progress from optimistic narrative.','Document assumptions and sensitivity.','Present ranges and confidence where appropriate.','Connect corrective actions to forecast movement.']),
 ('episode-02-contract-records','Notices, Records, and Protecting Contractual Position','Contract Management',
  ['Contract administration begins with understanding obligations, authority, time limits, and required communications.','A timely factual notice protects clarity and gives parties an opportunity to respond.','Good records are created during delivery and linked to decisions, events, instructions, cost, time, and performance.','Contract discipline and collaborative relationships can coexist when communication is accurate, respectful, and transparent.'],
  ['Maintain an obligation and notice register.','Use approved correspondence channels.','Separate fact, opinion, and reservation.','Link records to the relevant event and clause.','Escalate material deadlines before they expire.']),
 ('episode-03-sponsor-clarity','How Sponsors Create Clarity Before Delivery Pressure Increases','Governance',
  ['Sponsors connect the business outcome with authority, resources, priorities, risk appetite, and governance.','Teams need clarity about success, tolerances, decision rights, escalation, and benefits ownership.','Effective sponsorship challenges optimism while protecting constructive reporting.','Visible, timely decisions prevent unresolved ambiguity from becoming delivery delay.'],
  ['Confirm the outcome and measurable benefits.','Publish decision rights and tolerances.','Schedule stage gates around real decisions.','Remove organisational barriers promptly.','Review whether governance is helping delivery.']),
 ('episode-04-responsible-ai','Responsible AI for Project and Contract Decisions','Digital Practice',
  ['AI may support analysis, drafting, classification, forecasting, and knowledge retrieval, but accountability remains human.','Approved use cases, protected data, validation, explainability, and records are essential controls.','Confidential, personal, privileged, or contract-restricted information must not enter unapproved systems.','High-impact outputs require qualified review, challenge, and escalation.'],
  ['Define permitted and prohibited uses.','Classify data before using AI.','Validate outputs against reliable evidence.','Record material AI-assisted decisions.','Monitor bias, drift, error, and unintended consequences.']),
 ('episode-05-career-journey','From Project Coordinator to Programme Leader','Professional Growth',
  ['Career progress grows through reliable delivery, curiosity, feedback, broader responsibility, and evidence of outcomes.','Early roles build coordination, records, communication, planning, and control discipline.','Leadership roles require integration, judgement, governance, commercial awareness, and influence across uncertainty.','A useful development plan targets the next responsibility to practise rather than only the next title.'],
  ['Keep an evidence-based career portfolio.','Ask for specific feedback.','Seek stretch work with appropriate support.','Connect learning to real responsibilities.','Reflect on setbacks without hiding them.']),
 ('episode-06-baseline-recovery','What Recovery Teams Learn When the Baseline Is No Longer Credible','Delivery Recovery',
  ['Recovery begins by separating facts from assumptions and validating actual progress and remaining scope.','Uncontrolled change, weak productivity evidence, broken logic, and commercial uncertainty must be exposed.','The goal is not a more attractive completion date; it is a decision-useful control system.','Recovery also addresses behaviours, authority, resources, and governance that allowed unreliable information to persist.'],
  ['Establish a controlled data date.','Run a scope and change census.','Validate schedule logic and productivity.','Align project and contract positions.','Publish recovery decisions, owners, and evidence.']),
]

def set_cell_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr(); borders = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    for k,v in [('val','single'),('sz','12'),('space','5'),('color',RED)]: bottom.set(qn('w:'+k),v)
    borders.append(bottom); pPr.append(borders)

def configure_doc(doc, title):
    sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
    for name,size,before,after,color in [('Heading 1',16,18,10,NAVY),('Heading 2',13,14,7,NAVY),('Heading 3',12,10,5,'1F4D78')]:
        s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    header=sec.header.paragraphs[0]; header.text='PCMO Podcasts | Episode Transcript'; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs: r.font.name='Calibri'; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(SLATE)
    footer=sec.footer.paragraphs[0]; footer.text='Project & Contracts Management Organisation | Member learning resource'; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in footer.runs: r.font.name='Calibri'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(SLATE)

def add_docx(slug,title,category,points,takeaways):
    doc=Document(); configure_doc(doc,title)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); p.add_run('PCMO PODCASTS').bold=True; p.runs[0].font.color.rgb=RGBColor.from_string(RED); p.runs[0].font.size=Pt(12)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10); r=p.add_run(title); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(f'{category} | Full episode transcript and learning guide'); r.italic=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(SLATE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30); r=p.add_run('Project & Contracts Management Organisation'); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); set_cell_border(p)
    doc.add_page_break()
    doc.add_heading('Episode overview',1); doc.add_paragraph(f'This PCMO episode examines {title.lower()} through the lens of professional practice. It is designed to help listeners connect principles with decisions, controls, behaviours, and evidence in real project and contract environments.')
    doc.add_heading('Learning objectives',2)
    for item in ['Recognise the professional issue and why it matters.','Identify practical controls, questions, and behaviours.','Evaluate how context changes the appropriate response.','Select actions for further learning and workplace application.']: doc.add_paragraph(item,style='List Bullet')
    doc.add_heading('Full discussion transcript',1)
    for i,point in enumerate(points,1):
        doc.add_heading(f'{i}. {point.split(".")[0]}',2)
        doc.add_paragraph(point)
        doc.add_paragraph(f'In practice, this requires professionals to test the quality of available information, clarify ownership and authority, consider risk and contractual context, communicate uncertainty honestly, and retain evidence of material decisions. The correct response should be proportionate to scale, complexity, urgency, stakeholder impact, and applicable obligations.')
        doc.add_paragraph(f'The discussion also emphasises constructive challenge. Teams should ask what is known, what remains assumed, which evidence can be verified, what could change the conclusion, who is authorised to decide, and how the outcome will be monitored. These questions turn the topic from general advice into a controlled professional conversation.')
    doc.add_heading('Practical takeaways',1)
    for item in takeaways: doc.add_paragraph(item,style='List Bullet')
    doc.add_heading('Reflection questions',1)
    for item in ['Where does this issue appear in your current environment?','Which existing control is strongest, and which needs improvement?','What evidence would demonstrate a better decision or outcome?','Who should be involved before action is taken?','What will you do differently in the next thirty days?']: doc.add_paragraph(item,style='List Number')
    doc.add_heading('Related PCMO learning',1); doc.add_paragraph('Continue through the PCMO Learning, Standards, Certifications, Career Resources, Thought Leadership, and Library pages. Apply guidance proportionately and obtain qualified legal, contractual, financial, technical, or professional advice where required.')
    doc.add_heading('Professional-use notice',1); doc.add_paragraph('This transcript supports education and professional discussion. It does not replace governing law, contract terms, organisational policy, standards, or specialist advice. Protect confidential and personal information when discussing examples.')
    path=OUT_DOCX/f'{slug}.docx'; doc.save(path); copy2(path,PUBLIC/path.name); return path

def add_pdf(slug,title,category,points,takeaways):
    path=OUT_PDF/f'{slug}.pdf'; styles=getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CoverTitle',fontName='Helvetica-Bold',fontSize=26,leading=32,textColor=HexColor('#0B3764'),alignment=TA_CENTER,spaceAfter=16))
    styles.add(ParagraphStyle(name='Kicker',fontName='Helvetica-Bold',fontSize=11,textColor=HexColor('#DC2626'),alignment=TA_CENTER,spaceAfter=24))
    styles.add(ParagraphStyle(name='BodyX',fontName='Helvetica',fontSize=9.5,leading=13,textColor=HexColor('#475569'),spaceAfter=6))
    styles.add(ParagraphStyle(name='H1X',fontName='Helvetica-Bold',fontSize=16,leading=19,textColor=HexColor('#0B3764'),spaceBefore=9,spaceAfter=6))
    styles.add(ParagraphStyle(name='H2X',fontName='Helvetica-Bold',fontSize=12,leading=14,textColor=HexColor('#0B3764'),spaceBefore=6,spaceAfter=4))
    def pages(canvas,doc):
        canvas.saveState(); canvas.setFont('Helvetica',8); canvas.setFillColor(HexColor('#64748B')); canvas.drawString(inch,.55*inch,'PCMO Podcasts | Episode Transcript'); canvas.drawRightString(7.5*inch,.55*inch,f'Page {doc.page}'); canvas.restoreState()
    story=[Spacer(1,1.15*inch),Image(str(LOGO),width=5.3*inch,height=.82*inch),Spacer(1,.45*inch),Paragraph('PCMO PODCASTS',styles['Kicker']),Paragraph(title,styles['CoverTitle']),Paragraph(f'{category} | Full episode transcript and learning guide',ParagraphStyle(name='Sub',parent=styles['BodyX'],alignment=TA_CENTER,fontSize=12)),PageBreak(),Paragraph('Episode overview',styles['H1X']),Paragraph(f'This PCMO episode examines {title.lower()} through professional practice, decisions, controls, behaviours, and evidence.',styles['BodyX']),Paragraph('Learning objectives',styles['H2X'])]
    for item in ['Recognise why the issue matters.','Identify practical controls and behaviours.','Evaluate contextual factors.','Select actions for further application.']: story.append(Paragraph(f'- {item}',styles['BodyX']))
    story.append(Paragraph('Full discussion transcript',styles['H1X']))
    for i,point in enumerate(points,1):
        story += [Paragraph(f'{i}. {point.split(".")[0]}',styles['H2X']),Paragraph(point,styles['BodyX']),Paragraph('In practice, professionals should test information quality, clarify ownership and authority, consider risk and contractual context, communicate uncertainty honestly, and retain evidence of material decisions. The response should be proportionate to scale, complexity, urgency, stakeholder impact, and applicable obligations.',styles['BodyX']),Paragraph('Constructive challenge asks what is known, what remains assumed, which evidence is verifiable, what could change the conclusion, who can decide, and how the outcome will be monitored.',styles['BodyX'])]
    story.append(Paragraph('Practical takeaways',styles['H1X']))
    for item in takeaways: story.append(Paragraph(f'- {item}',styles['BodyX']))
    story.append(Paragraph('Reflection questions',styles['H1X']))
    for i,item in enumerate(['Where does this issue appear in your environment?','Which control needs improvement?','What evidence would demonstrate progress?','Who should be involved?','What will you do in the next thirty days?'],1): story.append(Paragraph(f'{i}. {item}',styles['BodyX']))
    story += [Paragraph('Related PCMO learning',styles['H1X']),Paragraph('Continue through PCMO Learning, Standards, Certifications, Career Resources, Thought Leadership, and Library resources.',styles['BodyX']),Paragraph('Professional-use notice',styles['H1X']),Paragraph('This transcript supports education and professional discussion. It does not replace law, contract terms, organisational policy, standards, or specialist advice.',styles['BodyX'])]
    SimpleDocTemplate(str(path),pagesize=LETTER,rightMargin=inch,leftMargin=inch,topMargin=.8*inch,bottomMargin=.8*inch,title=title,author='PCMO').build(story,onFirstPage=pages,onLaterPages=pages)
    copy2(path,PUBLIC/path.name); return path

for episode in EPISODES:
    print(add_docx(*episode)); print(add_pdf(*episode))
